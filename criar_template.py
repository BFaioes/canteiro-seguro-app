import docx

NOME_DO_TEMPLATE_WORD = "template_apr_gerado.docx"

doc = docx.Document()
doc.add_heading('ANÁLISE PRELIMINAR DE RISCO (APR)', level=1)

p = doc.add_paragraph()
p.add_run('Título: ').bold = True
doc.add_paragraph('{{ titulo_apr }}')

p = doc.add_paragraph()
p.add_run('Local: ').bold = True
doc.add_paragraph('{{ local }}')

p = doc.add_paragraph()
p.add_run('Data: ').bold = True
doc.add_paragraph('{{ data_elaboracao }}')

doc.add_heading('ETAPAS DA TAREFA, RISCOS E MEDIDAS DE CONTROLE', level=3)
table = doc.add_table(rows=2, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Etapa da Tarefa'
hdr[1].text = 'Perigos Identificados'
hdr[2].text = 'Riscos Associados'
hdr[3].text = 'Medidas de Controle Recomendadas'
hdr[4].text = 'Classificação do Risco Residual'

row = table.rows[1].cells
row[0].text = '{% for r in etapas_e_riscos %}{{ r.etapa_tarefa }}{% if not loop.last %}\n---\n{% endif %}{% endfor %}'
row[1].text = '{% for r in etapas_e_riscos %}{{ r.perigos_identificados|join("\\n- ") }}{% if not loop.last %}\n---\n{% endif %}{% endfor %}'
row[2].text = '{% for r in etapas_e_riscos %}{{ r.riscos_associados|join("\\n- ") }}{% if not loop.last %}\n---\n{% endif %}{% endfor %}'
row[3].text = '{% for r in etapas_e_riscos %}{{ r.medidas_de_controle_recomendadas|join("\\n- ") }}{% if not loop.last %}\n---\n{% endif %}{% endfor %}'
row[4].text = '{% for r in etapas_e_riscos %}{{ r.classificacao_risco_residual }}{% if not loop.last %}\n---\n{% endif %}{% endfor %}'

doc.add_heading('EQUIPAMENTOS DE PROTEÇÃO INDIVIDUAL (EPIs) OBRIGATÓRIOS', level=3)
doc.add_paragraph('{{ epis_obrigatorios|join("\\n- ") }}')

doc.add_heading('PROCEDIMENTOS DE EMERGÊNCIA', level=3)
doc.add_paragraph('{{ procedimentos_emergencia }}')

doc.save(NOME_DO_TEMPLATE_WORD)
print(f"✅ Template '{NOME_DO_TEMPLATE_WORD}' criado com sucesso.")
