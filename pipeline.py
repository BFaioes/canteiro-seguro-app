# ======================================================================================
# PROJETO CANTEIRO SEGURO (v4.0) - STREAMLIT ADAPTAÇÃO
# ======================================================================================

import streamlit as st
import vertexai
from vertexai.generative_models import GenerativeModel
from vertexai.language_models import TextEmbeddingModel
from google.cloud import storage
from sklearn.metrics.pairwise import cosine_similarity
from docxtpl import DocxTemplate
import docx
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import json
import numpy as np
from datetime import datetime
import pypdf
from langchain.text_splitter import RecursiveCharacterTextSplitter
import tempfile
import os
from google.oauth2 import service_account

# --------------------------------------------------------------------------------------
# CONFIGURAÇÃO INICIAL DO STREAMLIT
# --------------------------------------------------------------------------------------

st.set_page_config(
    page_title="Canteiro Seguro - APR Generator",
    page_icon="🛡️",
    layout="wide"
)

st.title("🛡️ Canteiro Seguro - Gerador de APR")
st.markdown("---")

# --------------------------------------------------------------------------------------
# ETAPA 1: CONFIGURAÇÕES E AUTENTICAÇÃO
# --------------------------------------------------------------------------------------

# Configurações do projeto
PROJECT_ID = "arctic-dynamo-467600-k9"
BUCKET_NAME = "documentos-apr-bruno-revisao01"
LOCATION = "us-central1"

# Modelos
MODELO_DE_EMBEDDING = "text-embedding-004"
MODELO_DE_GERACAO = "gemini-2.5-flash"
TOP_K_RAG = 3

# Nomes dos arquivos
NOME_DO_ARQUIVO_FINAL = "APR_FINALMENTE_PREENCHIDA.docx"

# Autenticação via service account (para Streamlit Cloud)
@st.cache_resource
def init_vertexai():
    try:
        # Para Streamlit Cloud, use variáveis de ambiente
        credentials = service_account.Credentials.from_service_account_info(
            st.secrets["gcp_service_account"]
        )
        vertexai.init(project=PROJECT_ID, location=LOCATION, credentials=credentials)
        return True
    except Exception as e:
        st.error(f"Erro na autenticação: {e}")
        return False

# --------------------------------------------------------------------------------------
# ETAPA 2: FUNÇÕES AUXILIARES
# --------------------------------------------------------------------------------------

def criar_template_word():
    """Cria o template Word para a APR"""
    try:
        doc = docx.Document()
        doc.add_heading('ANÁLISE PRELIMINAR DE RISCO (APR)', level=1)

        # Informações básicas
        p = doc.add_paragraph()
        p.add_run('Título: ').bold = True
        doc.add_paragraph('{{ titulo_apr }}')

        p = doc.add_paragraph()
        p.add_run('Local: ').bold = True
        doc.add_paragraph('{{ local }}')

        p = doc.add_paragraph()
        p.add_run('Data: ').bold = True
        doc.add_paragraph('{{ data_elaboracao }}')

        # Tabela de etapas e riscos
        doc.add_heading('ETAPAS DA TAREFA, RISCOS E MEDIDAS DE CONTROLE', level=3)
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Etapa da Tarefa'
        hdr[1].text = 'Perigos Identificados'
        hdr[2].text = 'Riscos Associados'
        hdr[3].text = 'Medidas de Controle Recomendadas'
        hdr[4].text = 'Classificação do Risco Residual'

        row = table.rows[1].cells
        row[0].text = '{% for r in etapas_e_riscos %}{{ r.etapa_tarefa }}{% if not loop.last %}\n---\n{% endif %}{% endfor %}'
        row[1].text = '{% for r in etapas_e_riscos %}{{ r.perigos_identificados|join("\\n- ") }}{% if not loop.last %}\n---\n{% endif %}{% endfor %}'
        row[2].text = '{% for r in etapas_e_riscos %}{{ r.riscos_associados|join("\\n- ") }}{% if not loop.last %}\n---\n{% endif %}{% endfor %}'
        row[3].text = '{% for r in etapas_e_riscos %}{{ r.medidas_de_controle_recomendadas|join("\\n- ") }}{% if not loop.last %}\n---\n{% endif %}{% endfor %}'
        row[4].text = '{% for r in etapas_e_riscos %}{{ r.classificacao_risco_residual }}{% if not loop.last %}\n---\n{% endif %}{% endfor %}'

        # EPIs e procedimentos
        doc.add_heading('EQUIPAMENTOS DE PROTEÇÃO INDIVIDUAL (EPIs) OBRIGATÓRIOS', level=3)
        doc.add_paragraph('{{ epis_obrigatorios|join("\\n- ") }}')

        doc.add_heading('PROCEDIMENTOS DE EMERGÊNCIA', level=3)
        doc.add_paragraph('{{ procedimentos_emergencia }}')

        return doc
    except Exception as e:
        st.error(f"Erro ao criar template: {e}")
        return None

def processar_pdfs():
    """Processa PDFs do bucket GCS"""
    try:
        credentials = service_account.Credentials.from_service_account_info(
            st.secrets["gcp_service_account"]
        )
        storage_client = storage.Client(credentials=credentials, project=PROJECT_ID)
        bucket = storage_client.bucket(BUCKET_NAME)
        all_chunks = []

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        MAX_CHARS = 4000

        progress_bar = st.progress(0)
        status_text = st.empty()

        blobs = list(bucket.list_blobs())
        total_blobs = len([b for b in blobs if b.name.lower().endswith(".pdf")])
        
        if total_blobs == 0:
            st.warning("Nenhum PDF encontrado no bucket.")
            return []

        processed = 0
        for blob in blobs:
            if blob.name.lower().endswith(".pdf"):
                try:
                    status_text.text(f"Processando: {blob.name}")
                    with blob.open("rb") as pdf_file:
                        pdf_reader = pypdf.PdfReader(pdf_file)
                        pdf_text = "".join(page.extract_text() or "" for page in pdf_reader.pages)
                        
                        if not pdf_text.strip():
                            st.warning(f"PDF sem texto: {blob.name}")
                            continue

                        chunks = text_splitter.split_text(pdf_text)
                        safe_chunks = [c[:MAX_CHARS] if len(c) > MAX_CHARS else c for c in chunks]
                        for chunk in safe_chunks:
                            all_chunks.append({"source": blob.name, "content": chunk})
                    
                    processed += 1
                    progress_bar.progress(processed / total_blobs)
                    
                except Exception as e:
                    st.warning(f"Erro ao processar {blob.name}: {e}")
                    continue

        status_text.text("✅ Processamento de PDFs concluído!")
        return all_chunks

    except Exception as e:
        st.error(f"Erro ao acessar bucket: {e}")
        return []

def gerar_embeddings(all_chunks):
    """Gera embeddings para os chunks de texto"""
    try:
        embedding_model = TextEmbeddingModel.from_pretrained(MODELO_DE_EMBEDDING)
        vectors, batch_texts, current_size = [], [], 0
        MAX_BATCH_CHARS = 15000

        progress_bar = st.progress(0)
        status_text = st.empty()

        total_chunks = len(all_chunks)
        
        for i, chunk in enumerate(all_chunks):
            text = chunk["content"]
            if current_size + len(text) > MAX_BATCH_CHARS and batch_texts:
                res = embedding_model.get_embeddings(batch_texts)
                vectors.extend([e.values for e in res])
                batch_texts, current_size = [], 0
            
            batch_texts.append(text)
            current_size += len(text)
            
            progress_bar.progress((i + 1) / total_chunks)
            status_text.text(f"Vetorizando chunk {i+1}/{total_chunks}")

        if batch_texts:
            res = embedding_model.get_embeddings(batch_texts)
            vectors.extend([e.values for e in res])

        status_text.text("✅ Vetorização concluída!")
        return np.array(vectors)

    except Exception as e:
        st.error(f"Erro na geração de embeddings: {e}")
        return None

def executar_rag(tarefa_usuario, all_chunks, embeddings_array):
    """Executa o processo RAG"""
    try:
        embedding_model = TextEmbeddingModel.from_pretrained(MODELO_DE_EMBEDDING)
        
        # Embedding da query
        query_embedding = embedding_model.get_embeddings([tarefa_usuario])[0].values
        
        # Similaridade
        similarities = cosine_similarity([query_embedding], embeddings_array)[0]
        top_indices = similarities.argsort()[-TOP_K_RAG:][::-1]
        selecionados = [all_chunks[i]['content'] for i in top_indices]

        # Resumir chunks
        gen = GenerativeModel(MODELO_DE_GERACAO)
        resumos = []
        
        for i, chunk in enumerate(selecionados):
            with st.spinner(f"Resumindo contexto {i+1}/{len(selecionados)}..."):
                resumo_prompt = f"Resuma tecnicamente este trecho em até 80 palavras, preservando riscos, medidas e citações de NRs:\n\n{chunk[:2000]}"
                resumo_resp = gen.generate_content(resumo_prompt)
                resumos.append(resumo_resp.text.strip()[:500])

        contexto_recuperado = "\n---\n".join(resumos)
        if len(contexto_recuperado) > 3000:
            contexto_recuperado = contexto_recuperado[:3000] + "... [texto truncado]"

        return contexto_recuperado

    except Exception as e:
        st.error(f"Erro no RAG: {e}")
        return ""

def gerar_apr_ia(tarefa_usuario, contexto_recuperado):
    """Gera a APR usando IA"""
    try:
        gen = GenerativeModel(MODELO_DE_GERACAO)
        data_str = datetime.now().strftime("%d/%m/%Y")
        
        json_exemplo = '''{
  "titulo_apr": "APR - ATIVIDADE",
  "local": "Local",
  "data_elaboracao": "DATA",
  "etapas_e_riscos": [
    {
      "etapa_tarefa": "Etapa",
      "perigos_identificados": ["Perigo1"],
      "riscos_associados": ["Risco1"],
      "medidas_de_controle_recomendadas": ["Medida1 - NR XX"],
      "classificacao_risco_residual": "Baixo/Médio/Alto"
    }
  ],
  "epis_obrigatorios": ["EPI1"],
  "procedimentos_emergencia": "Procedimento"
}'''

        prompt_final = f"""
# PERSONA
Você é um Engenheiro de Segurança do Trabalho Sênior, especialista nas NRs e em análise de riscos.

# CONTEXTO (resumos de normas)
{contexto_recuperado}

# ATIVIDADE
{tarefa_usuario}

# TAREFA
Preencha a APR em JSON no formato abaixo, respeitando as regras:

{json_exemplo}

# REGRAS
- Preencha TODOS os campos do JSON.
- "epis_obrigatorios" nunca pode ficar vazio.
- Cite sempre a NR correspondente nas medidas de controle.
- Responda SOMENTE com JSON válido.
"""

        with st.spinner("Gerando APR com IA..."):
            response = gen.generate_content(prompt_final)
            
            if not response.text:
                raise ValueError("Resposta vazia da IA")

            # Extrair JSON
            txt = response.text.strip()
            start, end = txt.find('{'), txt.rfind('}') + 1
            dados_da_apr = json.loads(txt[start:end])

            # Pós-processamento
            atividade = tarefa_usuario.lower()

            if not dados_da_apr.get("epis_obrigatorios"):
                dados_da_apr["epis_obrigatorios"] = ["Capacete", "Botas", "Luvas", "Óculos"]

            if "altura" in atividade or "andaime" in atividade:
                dados_da_apr["epis_obrigatorios"].append("Cinto de segurança tipo paraquedista com talabarte duplo")

            if "confinado" in atividade or "espaço confinado" in atividade:
                dados_da_apr["epis_obrigatorios"] += ["Detector de gases", "Máscara com filtro", "Cinturão de segurança"]

            if not dados_da_apr.get("procedimentos_emergencia"):
                dados_da_apr["procedimentos_emergencia"] = (
                    "Acionar brigada de emergência, aplicar primeiros socorros, evacuar a área e acionar o Corpo de Bombeiros (193)."
                )

            return dados_da_apr

    except Exception as e:
        st.error(f"Erro na geração da APR: {e}")
        return None

def criar_documento_final(dados_da_apr):
    """Cria o documento Word final"""
    try:
        doc = docx.Document()

        # Título centralizado
        title = doc.add_heading('ANÁLISE PRELIMINAR DE RISCO (APR)', level=1)
        title.alignment = 1

        # Informações básicas
        p = doc.add_paragraph(); p.add_run('Título: ').bold = True
        doc.add_paragraph(dados_da_apr.get("titulo_apr", "N/A"))

        p = doc.add_paragraph(); p.add_run('Local: ').bold = True
        doc.add_paragraph(dados_da_apr.get("local", "N/A"))

        p = doc.add_paragraph(); p.add_run('Data: ').bold = True
        doc.add_paragraph(dados_da_apr.get("data_elaboracao", datetime.now().strftime("%d/%m/%Y")))

        # Tabela de etapas e riscos
        doc.add_heading('ETAPAS DA TAREFA, RISCOS E MEDIDAS DE CONTROLE', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = [
            "Etapa da Tarefa", "Perigos Identificados", "Riscos Associados",
            "Medidas de Controle Recomendadas", "Classificação do Risco Residual"
        ]
        
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
            hdr[i].paragraphs[0].alignment = 1
            shading = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            hdr[i]._tc.get_or_add_tcPr().append(shading)

        for etapa in dados_da_apr.get("etapas_e_riscos", []):
            row = table.add_row().cells
            row[0].text = etapa.get("etapa_tarefa", "N/A")
            row[1].text = "\n".join(etapa.get("perigos_identificados", []))
            row[2].text = "\n".join(etapa.get("riscos_associados", []))
            row[3].text = "\n".join(etapa.get("medidas_de_controle_recomendadas", []))
            row[4].text = etapa.get("classificacao_risco_residual", "N/A")
            
            for cell in row:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(6)
                    paragraph.alignment = 0

        # EPIs
        doc.add_heading('EQUIPAMENTOS DE PROTEÇÃO INDIVIDUAL (EPIs) OBRIGATÓRIOS', level=2)
        for epi in dados_da_apr.get("epis_obrigatorios", []):
            doc.add_paragraph(f"- {epi}", style="List Bullet")

        # Procedimentos
        doc.add_heading('PROCEDIMENTOS DE EMERGÊNCIA', level=2)
        doc.add_paragraph(dados_da_apr.get("procedimentos_emergencia", "N/A"))

        return doc

    except Exception as e:
        st.error(f"Erro ao criar documento: {e}")
        return None

# --------------------------------------------------------------------------------------
# INTERFACE PRINCIPAL DO STREAMLIT
# --------------------------------------------------------------------------------------

def main():
    # Sidebar com informações
    st.sidebar.title("ℹ️ Informações")
    st.sidebar.markdown("""
    **Como usar:**
    1. Digite a atividade/serviço
    2. Clique em Gerar APR
    3. Aguarde o processamento
    4. Baixe o documento gerado
    
    **Funcionalidades:**
    - ✅ Análise de riscos automatizada
    - ✅ Baseada em normas técnicas
    - ✅ Formatação profissional
    - ✅ Download em Word
    """)

    # Input principal
    st.subheader("📋 Descreva a atividade/serviço")
    tarefa_usuario = st.text_area(
        "Digite a atividade para gerar a APR:",
        placeholder="Ex: Trabalho em altura com instalação de equipamentos em torre de 15 metros...",
        height=100
    )

    if st.button("🛡️ Gerar APR", type="primary", use_container_width=True):
        if not tarefa_usuario.strip():
            st.error("Por favor, descreva a atividade/serviço.")
            return

        # Inicializar Vertex AI
        if not init_vertexai():
            return

        # Processar em etapas
        with st.status("Processando...", expanded=True) as status:
            # Etapa 1: Processar PDFs
            status.update(label="📚 Processando PDFs do bucket...", state="running")
            all_chunks = processar_pdfs()
            
            if not all_chunks:
                st.error("Nenhum conteúdo válido encontrado nos PDFs.")
                return

            # Etapa 2: Gerar embeddings
            status.update(label="🔢 Gerando embeddings...", state="running")
            embeddings_array = gerar_embeddings(all_chunks)
            
            if embeddings_array is None:
                st.error("Erro na geração de embeddings.")
                return

            # Etapa 3: Executar RAG
            status.update(label="🔍 Buscando contexto relevante...", state="running")
            contexto_recuperado = executar_rag(tarefa_usuario, all_chunks, embeddings_array)

            # Etapa 4: Gerar APR com IA
            status.update(label="🤖 Gerando análise de riscos...", state="running")
            dados_da_apr = gerar_apr_ia(tarefa_usuario, contexto_recuperado)
            
            if dados_da_apr is None:
                st.error("Erro na geração da APR.")
                return

            # Etapa 5: Criar documento
            status.update(label="📄 Formatando documento...", state="running")
            doc_final = criar_documento_final(dados_da_apr)
            
            if doc_final is None:
                st.error("Erro na criação do documento.")
                return

            status.update(label="✅ Processo concluído!", state="complete")

        # Salvar e disponibilizar download
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                doc_final.save(tmp_file.name)
                
                with open(tmp_file.name, "rb") as file:
                    st.success("APR gerada com sucesso!")
                    st.download_button(
                        label="📥 Download da APR",
                        data=file,
                        file_name=NOME_DO_ARQUIVO_FINAL,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            # Limpar arquivo temporário
            os.unlink(tmp_file.name)

        except Exception as e:
            st.error(f"Erro ao salvar documento: {e}")

        # Preview dos dados
        with st.expander("📊 Visualizar dados gerados"):
            st.json(dados_da_apr)

if __name__ == "__main__":
    main()
