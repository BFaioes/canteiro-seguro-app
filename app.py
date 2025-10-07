# ======================================================================================
# PROJETO CANTEIRO SEGURO - VERSÃO STREAMLIT
# ======================================================================================

# --------------------------------------------------------------------------------------
# IMPORTAÇÕES E CONFIGURAÇÃO INICIAL
# --------------------------------------------------------------------------------------
import streamlit as st
import vertexai
from vertexai.generative_models import GenerativeModel
from vertexai.language_models import TextEmbeddingModel
from google.cloud import storage
from google.oauth2 import service_account
from sklearn.metrics.pairwise import cosine_similarity
import docx
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import json
import numpy as np
from datetime import datetime
import pypdf
from langchain.text_splitter import RecursiveCharacterTextSplitter
import io

# --------------------------------------------------------------------------------------
# CONFIGURAÇÃO DA PÁGINA STREAMLIT
# --------------------------------------------------------------------------------------
st.set_page_config(
    page_title="Gerador de APR com IA",
    page_icon="👷",
    layout="wide"
)

# --------------------------------------------------------------------------------------
# FUNÇÕES COM CACHE (PARA PERFORMANCE)
# --------------------------------------------------------------------------------------

# Inicializa a conexão com o Vertex AI (cacheado para não reconectar sempre)
@st.cache_resource
def inicializar_vertexai():
    """Inicializa e autentica no Google Cloud e Vertex AI."""
    try:
        creds_dict = st.secrets["gcp"]
        credentials = service_account.Credentials.from_service_account_info(creds_dict)
        vertexai.init(
            project=creds_dict["project_id"],
            location="us-central1",
            credentials=credentials
        )
        storage_client = storage.Client(
            project=creds_dict["project_id"],
            credentials=credentials
        )
        return storage_client
    except Exception as e:
        st.error(f"Erro na autenticação com o Google Cloud. Verifique seus secrets. Detalhe: {e}")
        return None

# Carrega e processa os PDFs (cacheado para não reprocessar os arquivos)
@st.cache_data(ttl=3600) # Cache por 1 hora
def carregar_e_processar_pdfs(_storage_client):
    """Baixa os PDFs do bucket, extrai o texto e divide em chunks."""
    if not _storage_client:
        return []
        
    bucket_name = st.secrets["gcp"]["bucket_name"]
    bucket = _storage_client.bucket(bucket_name)
    all_chunks = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=700, chunk_overlap=100)
    
    pdf_files = [blob for blob in bucket.list_blobs() if blob.name.lower().endswith(".pdf")]

    if not pdf_files:
        st.warning("Nenhum arquivo PDF encontrado no bucket do Google Cloud Storage.")
        return []

    progress_bar = st.progress(0, text="Processando arquivos PDF...")
    for i, blob in enumerate(pdf_files):
        try:
            with blob.open("rb") as pdf_file:
                pdf_reader = pypdf.PdfReader(pdf_file)
                pdf_text = "".join(page.extract_text() or "" for page in pdf_reader.pages)
                if pdf_text.strip():
                    chunks = text_splitter.split_text(pdf_text)
                    for chunk in chunks:
                        all_chunks.append({"source": blob.name, "content": chunk})
        except Exception as e:
            st.warning(f"Não foi possível processar o arquivo {blob.name}: {e}")
        progress_bar.progress((i + 1) / len(pdf_files), text=f"Processando: {blob.name}")
    
    progress_bar.empty()
    return all_chunks

# Gera os embeddings (vetores) para os chunks de texto
@st.cache_data(ttl=3600)
def gerar_embeddings(_chunks):
    """Gera embeddings para uma lista de textos usando um modelo da Vertex AI."""
    if not _chunks:
        return np.array([])
    
    model = TextEmbeddingModel.from_pretrained("text-embedding-004")
    text_contents = [chunk["content"] for chunk in _chunks]
    
    # Processa em lotes para evitar limites da API
    batch_size = 250 
    all_embeddings = []
    for i in range(0, len(text_contents), batch_size):
        batch = text_contents[i:i + batch_size]
        res = model.get_embeddings(batch)
        all_embeddings.extend([e.values for e in res])
        
    return np.array(all_embeddings)

# --------------------------------------------------------------------------------------
# FUNÇÃO PRINCIPAL DA LÓGICA DE NEGÓCIO
# --------------------------------------------------------------------------------------

def gerar_apr_completa(tarefa_do_usuario, chunks, embeddings_array):
    """Orquestra o processo de RAG e geração de conteúdo pela IA."""
    
    # 1. Similaridade (RAG)
    with st.spinner("Buscando informações relevantes nas normas..."):
        embedding_model = TextEmbeddingModel.from_pretrained("text-embedding-004")
        query_embedding = embedding_model.get_embeddings([tarefa_do_usuario])[0].values
        
        similarities = cosine_similarity([query_embedding], embeddings_array)[0]
        top_indices = similarities.argsort()[-3:][::-1] # TOP_K = 3
        
        contexto_chunks = [chunks[i]['content'] for i in top_indices]
        contexto_recuperado = "\n\n---\n\n".join(contexto_chunks)

    # 2. Geração com LLM (Gemini)
    with st.spinner("IA (Eng. de Segurança Sênior) está redigindo a APR..."):
        modelo_generativo = GenerativeModel("gemini-1.5-flash-001")
        
        # O JSON de exemplo é melhor deixar fora do prompt principal para clareza
        json_exemplo = """{
            "titulo_apr": "APR - Título da Atividade",
            "local": "A definir",
            "data_elaboracao": "DD/MM/AAAA",
            "etapas_e_riscos": [
                {
                    "etapa_tarefa": "Ex: Preparação da Área",
                    "perigos_identificados": ["Ex: Piso irregular"],
                    "riscos_associados": ["Ex: Queda de mesmo nível"],
                    "medidas_de_controle_recomendadas": ["Ex: Isolar e sinalizar a área - NR 18"],
                    "classificacao_risco_residual": "Baixo/Médio/Alto"
                }
            ],
            "epis_obrigatorios": ["Capacete", "Botas de segurança"],
            "procedimentos_emergencia": "Acionar brigada de emergência (ramal XXX), prestar primeiros socorros e ligar para emergência (192/193)."
        }"""

        prompt_final = f"""
        # PERSONA
        Você é um Engenheiro de Segurança do Trabalho Sênior, especialista em Normas Regulamentadoras (NRs) brasileiras e em análise de riscos para a construção civil. Sua linguagem é técnica, direta e focada na prevenção.

        # CONTEXTO TÉCNICO EXTRAÍDO DE NORMAS:
        {contexto_recuperado}

        # ATIVIDADE A SER ANALISADA:
        {tarefa_do_usuario}

        # TAREFA
        Com base no CONTEXTO TÉCNICO e em seu conhecimento especializado, preencha uma Análise Preliminar de Risco (APR) para a ATIVIDADE. A resposta deve ser um único e válido objeto JSON, seguindo estritamente o formato do exemplo abaixo.

        # FORMATO JSON OBRIGATÓRIO:
        {json_exemplo}

        # REGRAS CRÍTICAS:
        - Responda APENAS com o código JSON. Não inclua texto, explicações ou marcadores como ```json.
        - Preencha todos os campos do JSON com informações detalhadas e aplicáveis.
        - Nas "medidas_de_controle_recomendadas", sempre que possível, cite a NR correspondente (ex: "Instalar guarda-corpo de 1.20m - NR 18").
        - A "classificacao_risco_residual" deve ser "Alto" para atividades como trabalho em altura, espaços confinados, ou com inflamáveis. Para atividades com máquinas ou eletricidade, use "Médio". Use "Baixo" apenas para tarefas administrativas.
        - Os "epis_obrigatorios" e "procedimentos_emergencia" não podem ser vazios.
        """
        
        response = modelo_generativo.generate_content(prompt_final)
        
        try:
            # Limpa a resposta para garantir que seja um JSON válido
            json_text = response.text.strip().replace("```json", "").replace("```", "")
            dados_da_apr = json.loads(json_text)
        except (json.JSONDecodeError, AttributeError) as e:
            st.error(f"A IA retornou um formato inesperado. Tentando novamente... Detalhe do erro: {e}")
            st.code(response.text) # Mostra o que a IA retornou para depuração
            return None

    # 3. Geração do Documento Word
    with st.spinner("Formatando o documento Word..."):
        doc = docx.Document()
        doc.add_heading('ANÁLISE PRELIMINAR DE RISCO (APR)', level=1).alignment = 1

        doc.add_paragraph().add_run('Título: ').bold = True
        doc.add_paragraph(dados_da_apr.get("titulo_apr", tarefa_do_usuario))
        
        doc.add_paragraph().add_run('Local: ').bold = True
        doc.add_paragraph(dados_da_apr.get("local", "N/A"))
        
        doc.add_paragraph().add_run('Data: ').bold = True
        doc.add_paragraph(datetime.now().strftime("%d/%m/%Y"))

        doc.add_heading('ETAPAS DA TAREFA, RISCOS E MEDIDAS DE CONTROLE', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ["Etapa da Tarefa", "Perigos Identificados", "Riscos Associados", "Medidas de Controle Recomendadas", "Risco Residual"]
        
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = header_text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = 1
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        for etapa in dados_da_apr.get("etapas_e_riscos", []):
            row_cells = table.add_row().cells
            row_cells[0].text = etapa.get("etapa_tarefa", "")
            row_cells[1].text = "\n".join(f"- {p}" for p in etapa.get("perigos_identificados", []))
            row_cells[2].text = "\n".join(f"- {r}" for r in etapa.get("riscos_associados", []))
            row_cells[3].text = "\n".join(f"- {m}" for m in etapa.get("medidas_de_controle_recomendadas", []))
            row_cells[4].text = etapa.get("classificacao_risco_residual", "N/A")
            row_cells[4].paragraphs[0].alignment = 1

        doc.add_heading('EQUIPAMENTOS DE PROTEÇÃO INDIVIDUAL (EPIs)', level=2)
        for epi in dados_da_apr.get("epis_obrigatorios", []):
            doc.add_paragraph(epi, style='List Bullet')

        doc.add_heading('PROCEDIMENTOS DE EMERGÊNCIA', level=2)
        doc.add_paragraph(dados_da_apr.get("procedimentos_emergencia", "N/A"))

        # Salva o documento em um buffer de memória
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io

# --------------------------------------------------------------------------------------
# INTERFACE DO USUÁRIO (Layout do App)
# --------------------------------------------------------------------------------------

st.title("👷 Gerador de Análise Preliminar de Risco (APR)")
st.markdown("Desenvolvido com IA do Google (Vertex AI) e RAG para consulta em Normas Regulamentadoras.")

# Área de input do usuário (primeiro)
tarefa_usuario = st.text_area(
    "**Descreva a atividade ou serviço para a qual a APR será gerada:**",
    height=100,
    placeholder="Exemplo: Montagem de andaime fachadeiro com 15 metros de altura para reboco externo."
)

if st.button("Gerar APR", type="primary", use_container_width=True):
    if not tarefa_usuario:
        st.warning("Por favor, descreva a atividade antes de gerar a APR.")
    else:
        # Inicializa serviços e carrega base de conhecimento somente após o clique
        storage_client = inicializar_vertexai()
        if not storage_client:
            st.error("Falha na autenticação com o Google Cloud. Verifique os secrets.")
        else:
            chunks_de_texto = carregar_e_processar_pdfs(storage_client)
            if not chunks_de_texto:
                st.warning("Nenhum conteúdo disponível no bucket para consulta.")
            else:
                vetores = gerar_embeddings(chunks_de_texto)
                if len(vetores) == 0:
                    st.warning("Não foi possível gerar embeddings para a base de conhecimento.")
                else:
                    st.success(f"Base de conhecimento carregada com {len(chunks_de_texto)} trechos de normas.")
                    documento_word = gerar_apr_completa(tarefa_usuario, chunks_de_texto, vetores)
                    if documento_word:
                        st.balloons()
                        st.download_button(
                            label="✔️ Download da APR em .docx",
                            data=documento_word,
                            file_name=f"APR_{tarefa_usuario[:20].replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )